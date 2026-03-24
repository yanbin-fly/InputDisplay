# -*- coding: utf-8 -*-
"""
为 Python 中级开发者生成 InputDisplay 项目解读 Word 文档
依赖: pip install python-docx
运行: python generate_word_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 全局样式设置 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

HEADING_COLOR = RGBColor(0x20, 0x4F, 0x9E)  # 深蓝色

def set_run_font(run, size_pt=11, bold=False, color=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, size=18, color=None, space_before=18, space_after=6):
    p = doc.add_heading('', level=level)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = HEADING_COLOR
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_code_block(doc, code, language=''):
    """添加带底色的代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_normal_para(doc, text, bold=False, indent=0, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.bold = bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    return p

def add_divider(doc):
    p = doc.add_paragraph('─' * 60)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('InputDisplay 项目完整解读')
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('面向 Python 中级开发者的 C# / WPF 技术入门指南')
sub_run.font.size = Pt(13)
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_p.add_run(f'文档版本：v1.0.0    生成日期：{datetime.date.today()}')
info_run.font.size = Pt(10)
info_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 目录
# ════════════════════════════════════════════════════════════
add_heading(doc, '目 录', 1, 18)
contents = [
    '一、写给 Python 开发者的前言',
    '二、项目是什么，能做什么',
    '三、技术选型：为什么是 WPF',
    '四、开发环境准备',
    '五、核心概念快速上手（C# 篇）',
    '六、Win32 API 调用：Python 的 ctypes vs C# 的 P/Invoke',
    '七、全局钩子原理：从 Python 角度理解',
    '八、代码全览与解读',
    '九、WPF 数据绑定：MVVM 思想',
    '十、线程模型：协程 vs Dispatcher',
    '十一、编译、打包与发布',
    '十二、学习路径与资源推荐',
]
for i, item in enumerate(contents, 1):
    p = doc.add_paragraph()
    p.add_run(f'  {item}')
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 第一章
# ════════════════════════════════════════════════════════════
add_heading(doc, '一、写给 Python 开发者的前言', 1)

add_normal_para(doc, (
    '你好！这篇文档的假设读者是一位具备 Python 中级水平的开发者——'
    '你熟悉 Python 的基本语法、常用标准库（如 os、sys、threading）、'
    '以及 Python 中的一些高级概念（如装饰器、上下文管理器、类型注解）。'
    '但对于 C# 这门语言，以及 .NET 生态、Windows GUI 开发，你可能还比较陌生。'
))
add_normal_para(doc, (
    '这篇文档的目标，不是让你成为 C# 专家，而是帮你快速建立起足够的知识框架，'
    '让你能够理解这个项目的代码逻辑，甚至在此基础上做自己的修改和扩展。'
))
add_normal_para(doc, '我们会从 Python 的视角出发，逐一找到对应的 C# 概念，帮助你建立知识迁移。')

add_heading(doc, '1.1 先问一个问题：Python 能做这个吗？', 2, 14)

add_normal_para(doc, '答案是：可以，但各有各的难处。')
add_normal_para(doc, '如果你用 Python 来实现类似功能，可能的选择有：')
add_bullet(doc, 'pynput：轻量级，但需要额外依赖')
add_bullet(doc, 'pyHook / pywinhook：较老，部分 Python 版本不兼容')
add_bullet(doc, 'ctypes + Win32 API：可行，但 UI 部分（浮窗、半透明）需要额外方案（如 tkinter、PyQt）')
add_normal_para(doc, (
    '这个项目选择 C# + WPF，主要是因为：'
    'C# 对 Windows API 的互操作性是原生级别的，几乎没有额外开销；'
    'WPF 自带声明式 UI 框架，数据绑定非常强大；'
    '.NET 8.0 编译出的 exe 文件非常小（几十 KB），无需安装运行时也可以发布。'
))

# ════════════════════════════════════════════════════════════
# 第二章
# ════════════════════════════════════════════════════════════
add_heading(doc, '二、项目是什么，能做什么', 1)

add_heading(doc, '2.1 功能概览', 2, 14)
add_normal_para(doc, 'InputDisplay 是一款运行在 Windows 上的输入监控工具：')
add_bullet(doc, '实时捕获键盘按键：包括普通键、功能键、组合键（Ctrl+C、Win+D 等）')
add_bullet(doc, '实时捕获鼠标操作：点击、滚轮、坐标')
add_bullet(doc, '顶部卡片实时展示当前输入内容')
add_bullet(doc, '展开统计面板，查看按键频率 Top5')
add_bullet(doc, '浅色半透明浮窗，悬停在其他应用上方')
add_bullet(doc, '关闭按钮仅隐藏窗口，后台持续监控')

add_heading(doc, '2.2 技术亮点', 2, 14)
add_bullet(doc, '全局键盘/鼠标钩子：SetWindowsHookEx + Low-Level Hook')
add_bullet(doc, '声明式 UI：XAML 定义浮窗样式（圆角、半透明、动画）')
add_bullet(doc, '数据绑定：ObservableCollection 自动更新列表，INPC 驱动 UI 刷新')
add_bullet(doc, '系统托盘：WinForms NotifyIcon（5 行代码实现托盘集成）')
add_bullet(doc, '零依赖 UI 框架：无需第三方 UI 库，纯 WPF 内置功能')

# ════════════════════════════════════════════════════════════
# 第三章
# ════════════════════════════════════════════════════════════
add_heading(doc, '三、技术选型：为什么是 WPF', 1)

add_heading(doc, '3.1 .NET 生态 vs Python 生态', 2, 14)
add_normal_para(doc, '在开始之前，先看看技术栈的类比：')
add_code_block(doc, '''\
┌──────────────┬──────────────────┬──────────────────────┐
│    角色      │     Python       │        C# / WPF      │
├──────────────┼──────────────────┼──────────────────────┤
│  编程语言    │     Python       │         C#           │
│  版本        │     3.12         │        12            │
│  运行时      │     CPython      │      .NET CLR        │
│  GUI 框架    │  tkinter / PyQt  │        WPF           │
│  GUI 声明式  │    tkinter.ttk   │       XAML           │
│  数据绑定    │  PyQt signals    │   WPF Binding        │
│  包管理      │      pip         │       NuGet          │
│  编译目标    │   .pyc (字节码)   │   .exe / .dll        │
│  打包工具    │  pyinstaller     │    dotnet publish    │
└──────────────┴──────────────────┴──────────────────────┘''')

add_heading(doc, '3.2 WPF 是什么', 2, 14)
add_normal_para(doc, 'WPF（Windows Presentation Foundation）是微软 2006 年推出的 GUI 框架。相比 tkinter，它的核心优势是：')
add_bullet(doc, 'XAML 声明式 UI：UI 代码和逻辑代码分离，类似 HTML+CSS 和 JavaScript 的关系')
add_bullet(doc, '强大的数据绑定：数据和 UI 自动同步，几乎不需要手动更新界面')
add_bullet(doc, '矢量图形：支持圆角、渐变、阴影、动画等高级视觉效果')
add_bullet(doc, 'DirectX 渲染：界面渲染效率远高于基于 Canvas 的 tkinter')
add_normal_para(doc, '当然 WPF 也有局限：只能在 Windows 上运行（这也是我们项目的目标平台）。')

# ════════════════════════════════════════════════════════════
# 第四章
# ════════════════════════════════════════════════════════════
add_heading(doc, '四、开发环境准备', 1)

add_heading(doc, '4.1 安装 .NET SDK', 2, 14)
add_normal_para(doc, '从 https://dotnet.microsoft.com/download/dotnet/8.0 下载 .NET 8.0 SDK（软件开发者工具包）。')
add_normal_para(doc, '安装完成后，命令行验证：')
add_code_block(doc, '''\
C:\\> dotnet --version
8.0.419''')

add_heading(doc, '4.2 创建 WPF 项目（对比 Python 的venv）', 2, 14)
add_normal_para(doc, 'Python 创建虚拟环境：')
add_code_block(doc, '''\
python -m venv venv
venv\\Scripts\\activate  # Windows
pip install xxx''')

add_normal_para(doc, 'C# / .NET 创建项目（更简单，不需要手动管理虚拟环境）：')
add_code_block(doc, '''\
# 创建 WPF 项目（一条命令搞定，比 pip install + venv 还快）
dotnet new wpf -n InputDisplay -o InputDisplay

# 目录结构自动生成：
# InputDisplay/
#   ├── InputDisplay.csproj   ← 类似 requirements.txt + pyproject.toml
#   ├── App.xaml              ← 应用入口配置
#   ├── App.xaml.cs           ← 应用入口代码
#   ├── MainWindow.xaml        ← 主窗口 UI（类似 .kv 文件）
#   └── MainWindow.xaml.cs    ← 主窗口逻辑（类似 .py 文件）
''')

add_heading(doc, '4.3 项目配置（InputDisplay.csproj）', 2, 14)
add_normal_para(doc, '.csproj 文件相当于 Python 项目的 pyproject.toml 或 setup.py：')
add_code_block(doc, '''\
<Project Sdk="Microsoft.NET.Sdk">
  <PropertyGroup>
    <OutputType>WinExe</OutputType>
    <TargetFramework>net8.0-windows</TargetFramework>
    <UseWPF>true</UseWPF>
    <UseWindowsForms>true</UseWindowsForms>  ← 启用 WinForms（托盘图标需要）
    <AllowUnsafeBlocks>true</AllowUnsafeBlocks>
    <AssemblyName>InputDisplay</AssemblyName>
  </PropertyGroup>
</Project>''')
add_normal_para(doc, '对比 Python：不需要 requirements.txt，NuGet 包直接在 .csproj 里声明，dotnet restore 自动处理。')

# ════════════════════════════════════════════════════════════
# 第五章
# ════════════════════════════════════════════════════════════
add_heading(doc, '五、核心概念快速上手（C# 篇）', 1)

add_heading(doc, '5.1 命名空间 vs Python 模块', 2, 14)
add_normal_para(doc, 'Python 的 import 相当于 C# 的 using：')
add_code_block(doc, '''\
# Python
from collections import defaultdict
import os.path as path

# C# (写在文件顶部)
using System.Collections.ObjectModel;
using System.ComponentModel;''')

add_heading(doc, '5.2 类定义与继承（对比 Python）', 2, 14)
add_code_block(doc, '''\
# Python
class MyClass(BaseClass):
    def __init__(self, name):
        super().__init__(name)
        self._value = 0

    @property
    def value(self):
        return self._value

    def do_something(self):
        print("done")
''')

add_code_block(doc, '''\
// C#
public class MyClass : BaseClass, INotifyPropertyChanged  // 可多继承接口
{
    private int _value;  // 字段（类似 Python 的 self._xxx）

    // 属性（类似 Python 的 @property）
    public int Value
    {
        get => _value;
        set
        {
            _value = value;
            OnPropertyChanged();  // 通知 UI 刷新
        }
    }

    public void DoSomething()
    {
        Console.WriteLine("done");
    }
}''')

add_heading(doc, '5.3 事件与回调（对比 Python 装饰器）', 2, 14)
add_normal_para(doc, 'Python 中的事件绑定：')
add_code_block(doc, '''\
# Python ( tkinter 风格)
button.bind("<Button-1>", on_click)
button.bind("<KeyPress>", on_keypress)

def on_click(event):
    print(f"Clicked at {event.x}, {event.y}")
''')

add_normal_para(doc, 'C# 中的事件订阅（类似但更类型安全）：')
add_code_block(doc, '''\
// C# WPF 风格
button.Click += OnButtonClick;
listBox.MouseDoubleClick += OnListDoubleClick;

private void OnButtonClick(object sender, RoutedEventArgs e)
{
    // sender = 触发事件的控件
    // e = 事件参数
    Debug.WriteLine("Button clicked!");
}

private void OnListDoubleClick(object sender, MouseButtonEventArgs e)
{
    var point = e.GetPosition(listBox);
    Debug.WriteLine($"Double click at {point.X}, {point.Y}");
}
''')

add_heading(doc, '5.4 类型注解 vs var 推断', 2, 14)
add_code_block(doc, '''\
# Python
def process_input(event: InputEvent, stats: SessionStatistics) -> None:
    recent: list[InputEvent] = []
    top_keys: dict[str, int] = {}
    pass

# C# 两种写法都支持
void ProcessInput(InputEvent evt, SessionStatistics stats)
{
    // 显式类型（推荐阅读）
    ObservableCollection<InputEvent> recent = new();

    // var 推断（类似 Python 的动态类型）
    var topKeys = new Dictionary<string, int>();
}
''')

add_heading(doc, '5.5 lambda 表达式', 2, 14)
add_code_block(doc, '''\
# Python lambda
stats_timer.tick += lambda s, e: update_ui()

# C# lambda
_statsTimer.Tick += (_, _) => UpdateStats();
_button.Click += (s, e) => { Debug.WriteLine("Clicked!"); };
''')

add_heading(doc, '5.6 异步与主线程（对比 asyncio）', 2, 14)
add_normal_para(doc, 'Python asyncio 的协程调度：')
add_code_block(doc, '''\
# Python
async def fetch_data(url):
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as resp:
            return await resp.json()
''')

add_normal_para(doc, 'WPF 中的主线程调度（类似 asyncio 的 loop.call_soon_threadsafe）：')
add_code_block(doc, '''\
// C# - WPF 不需要 async/await 来做 UI 线程调度
// 只需要 Dispatcher.Invoke() 将代码调度到 UI 线程

// 从钩子回调线程（系统线程）调度到 UI 线程
Application.Current?.Dispatcher.Invoke(() =>
{
    CurrentInput = inputEvent;         // 触发 INPC，刷新绑定
    RecentInputs.Insert(0, inputEvent); // ObservableCollection 自动通知 UI
    Statistics.TotalKeyPresses++;        // 触发 INPC，刷新统计数字
});

// Dispatcher.Invoke() 类似 Python 的：
// loop = asyncio.get_event_loop()
// loop.call_soon_threadsafe(lambda: update_ui(), loop)
''')

# ════════════════════════════════════════════════════════════
# 第六章
# ════════════════════════════════════════════════════════════
add_heading(doc, '六、Win32 API 调用：Python 的 ctypes vs C# 的 P/Invoke', 1)

add_heading(doc, '6.1 Python ctypes 调用 Win32 API', 2, 14)
add_normal_para(doc, 'Python 中调用 Windows API 通常用 ctypes：')
add_code_block(doc, '''\
import ctypes
from ctypes import wintypes

# 定义函数签名
user32 = ctypes.windll.user32

# C 原型：BOOL SetWindowsHookExW(int idHook, HOOKPROC lpfn,
#                              HINSTANCE hmod, DWORD dwThreadId);
user32.SetWindowsHookExW.argtypes = [wintypes.INT, wintypes.HANDLE,
                                      wintypes.HINSTANCE, wintypes.DWORD]
user32.SetWindowsHookExW.restype = wintypes.HANDLE

# 调用
hook_id = user32.SetWindowsHookExW(13, callback_func, None, 0)
# 13 = WH_KEYBOARD_LL
''')

add_heading(doc, '6.2 C# P/Invoke 调用 Win32 API', 2, 14)
add_normal_para(doc, 'C# 的 P/Invoke 比 ctypes 更简洁，类型安全：')
add_code_block(doc, '''\
// C# - WinApi.cs
using System.Runtime.InteropServices;

public static class WinApi
{
    // C# 声明远比 Python 简洁，不需要手动指定类型映射
    [DllImport("user32.dll", CharSet = CharSet.Auto, SetLastError = true)]
    public static extern IntPtr SetWindowsHookEx(
        int idHook,
        LowLevelKeyboardProc lpfn,   // 委托类型，自动处理
        IntPtr hMod,
        uint dwThreadId
    );

    // 结构体自动映射（LayoutKind.Sequential = pack)
    [StructLayout(LayoutKind.Sequential)]
    public struct KBDLLHOOKSTRUCT
    {
        public uint VkCode;
        public uint ScanCode;
        public uint Flags;
        public uint Time;
        public IntPtr DwExtraInfo;
    }
}

// 调用
var hookId = WinApi.SetWindowsHookEx(
    WinApi.WH_KEYBOARD_LL,
    _keyboardProc,
    moduleHandle,
    0
);
''')

add_heading(doc, '6.3 对比总结', 2, 14)
add_code_block(doc, '''\
┌────────────────┬────────────────────┬──────────────────────────────┐
│    方面        │   Python ctypes     │      C# P/Invoke            │
├────────────────┼────────────────────┼──────────────────────────────┤
│  声明语法      │  分散在函数调用前    │  [DllImport] 属性装饰符     │
│  类型映射      │  手动 argtypes       │  自动（大部分）             │
│  结构体        │  ctypes.Structure   │  [StructLayout] 类          │
│  字符串编码    │  需要手动指定        │  CharSet 参数               │
│  错误处理      │  返回码检查          │  SetLastError + Exception  │
│  调用约定      │  默认 cdecl         │  默认 PASCAL（大部分 WinAPI）│
│  代码提示      │  无（运行时才发现）  │  有（编译期检查）           │
└────────────────┴────────────────────┴──────────────────────────────┘''')

# ════════════════════════════════════════════════════════════
# 第七章
# ════════════════════════════════════════════════════════════
add_heading(doc, '七、全局钩子原理：从 Python 角度理解', 1)

add_heading(doc, '7.1 什么是全局钩子（Hook）', 2, 14)
add_normal_para(doc, '全局钩子是 Windows 系统提供的一种机制，允许应用程序注册一个回调函数，'
    '当特定类型的系统事件（键盘、鼠标、键盘等）发生时，系统会调用这个回调函数。')
add_normal_para(doc, '类似于 Python 中的：')
add_code_block(doc, '''\
# Python 的事件监听器（但这只是应用层）
keyboard.on_press_key('a', lambda k: print("按了 a"))
mouse.on_click(lambda x, y: print(f"点击 ({x}, {y})"))

# 全局钩子相当于系统级的"超级监听器"——能捕获所有进程的所有输入
''')

add_heading(doc, '7.2 钩子链（Hook Chain）', 2, 14)
add_normal_para(doc, '系统中可能有多个程序注册了同类型的钩子，它们组成一个「钩子链」。'
    '事件发生时，链上的每个回调都会收到通知。你可以选择：')
add_bullet(doc, '处理事件（但不调用 CallNextHookEx）→ 阻止后续程序收到事件（不推荐）')
add_bullet(doc, '处理事件后调用 CallNextHookEx → 将事件传递给链上的下一个回调（推荐）')

add_heading(doc, '7.3 Low-Level Hook vs 普通 Hook', 2, 14)
add_normal_para(doc, '项目中用的是 Low-Level Hook（低级钩子），标记为 `_LL`：')
add_bullet(doc, 'WH_KEYBOARD_LL：低级键盘钩子，接收原始的键盘事件')
add_bullet(doc, 'WH_MOUSE_LL：低级鼠标钩子，接收原始的鼠标事件')
add_normal_para(doc, 'Low-Level Hook 的特殊之处：可以在托管代码（C#）中直接注册，不需要单独的 DLL。'
    '普通钩子则需要 DLL 注入到其他进程。')

add_heading(doc, '7.4 钩子回调流程图', 2, 14)
add_code_block(doc, '''\
Python 开发者可以这样类比理解：

┌──────────────────────────────────────────────────────────┐
│                    Windows 系统                          │
│                                                          │
│   键盘按下 ──► 消息队列 ──► 你的钩子回调 ──► 下一个钩子   │
│                          │                               │
│                    CallNextHookEx()                      │
│                          │                               │
└──────────────────────────┼───────────────────────────────┘
                           ▼
              ┌────────────────────────────┐
              │  Dispatcher.Invoke()       │  ← 将数据发送到 UI 线程
              │  (类似 asyncio 的         │
              │   call_soon_threadsafe)    │
              └────────────┬───────────────┘
                           ▼
              ┌────────────────────────────┐
              │  ObservableCollection     │  ← 类似 Python 的 notify
              │  .Insert(0, event)        │
              └────────────────────────────┘
                           ▼
              ┌────────────────────────────┐
              │  ListBox 自动渲染新行       │  ← UI 自动刷新
              └────────────────────────────┘
''')

# ════════════════════════════════════════════════════════════
# 第八章
# ════════════════════════════════════════════════════════════
add_heading(doc, '八、代码全览与解读', 1)

add_heading(doc, '8.1 WinApi.cs — Windows API 声明', 2, 14)
add_normal_para(doc, '这是整个项目的基础，声明了所有与 Windows 系统交互所需的函数和结构体：')
add_code_block(doc, '''\
public static class WinApi
{
    // 注册全局钩子（最核心的 API）
    [DllImport("user32.dll")]
    public static extern IntPtr SetWindowsHookEx(
        int idHook,        // 钩子类型：13=键盘，14=鼠标
        LowLevelKeyboardProc lpfn,  // 回调委托
        IntPtr hMod,       // 模块句柄
        uint dwThreadId    // 0=全局，所有进程
    );

    // 调用链上的下一个钩子（必须调用）
    [DllImport("user32.dll")]
    public static extern IntPtr CallNextHookEx(IntPtr hhk,
        int nCode, IntPtr wParam, IntPtr lParam);

    // 获取当前进程的模块句柄（传给 SetWindowsHookEx）
    [DllImport("kernel32.dll")]
    public static extern IntPtr GetModuleHandle(string? lpModuleName);

    // 判断某个虚拟键是否当前处于按下状态
    [DllImport("user32.dll")]
    public static extern short GetAsyncKeyState(int vKey);
}''')
add_normal_para(doc, '对比 Python：相当于把 ctypes 的所有 wintypes 定义和函数签名集中写在一个模块里。')

add_heading(doc, '8.2 Models.cs — 数据模型（对比 Pydantic）', 2, 14)
add_normal_para(doc, 'Python 中定义数据模型（类似 Pydantic）：')
add_code_block(doc, '''\
from pydantic import BaseModel
from datetime import datetime
from typing import Literal

class InputEvent(BaseModel):
    time: datetime = datetime.now
    type: Literal["Keyboard", "Mouse"]
    description: str = ""
    category: str = ""

class SessionStatistics(BaseModel):
    total_key_presses: int = 0
    total_mouse_clicks: int = 0
    key_frequency: dict[str, int] = {}
''')

add_normal_para(doc, 'C# 中对应实现（更啰嗦但更类型安全）：')
add_code_block(doc, '''\
// C# - Models.cs
public class InputEvent : INotifyPropertyChanged  // 实现通知接口
{
    private InputType _type;
    public InputType Type
    {
        get => _type;
        set { _type = value; OnPropertyChanged(); }  // 值变化时通知 UI
    }

    public string Description { get; set; } = "";

    // INotifyPropertyChanged 实现
    public event PropertyChangedEventHandler? PropertyChanged;
    protected void OnPropertyChanged([CallerMemberName] string? name = null)
        => PropertyChanged?.Invoke(this, new PropertyChangedEventArgs(name));
}

public enum InputType { Keyboard, Mouse }

public class SessionStatistics : INotifyPropertyChanged
{
    public int TotalKeyPresses { get; set; }
    // ...
}
''')
add_normal_para(doc, 'INotifyPropertyChanged（INPC）是 WPF 数据绑定的核心——Python 的 tkinter 没有这个机制，所以需要手动 update()。')

add_heading(doc, '8.3 InputHookService.cs — 钩子核心逻辑', 2, 14)
add_normal_para(doc, 'Python 中用 pynput 做键盘钩子：')
add_code_block(doc, '''\
from pynput import keyboard

def on_press(key):
    try:
        print(f"按了 {key.char}")
    except AttributeError:
        print(f"按了 {key}")

listener = keyboard.Listener(on_press=on_press)
listener.start()
listener.join()''')

add_normal_para(doc, 'C# 中用 Win32 API 做钩子（更底层但更强大）：')
add_code_block(doc, '''\
// 注册钩子（在 Start() 方法中）
var hookId = WinApi.SetWindowsHookEx(
    WinApi.WH_KEYBOARD_LL,
    _keyboardProc,     // 委托实例（保持 GC 不回收）
    moduleHandle,      // GetModuleHandle() 返回
    0                  // 全局
);

// 回调委托签名（必须匹配 Windows 期望的签名）
private IntPtr KeyboardHookCallback(int nCode, IntPtr wParam, IntPtr lParam)
{
    if (nCode >= 0)
    {
        // 从 lParam 解析键盘数据结构
        var hookStruct = Marshal.PtrToStructure<WinApi.KBDLLHOOKSTRUCT>(lParam);
        int vkCode = (int)hookStruct.VkCode;

        // 获取当前修饰键状态
        bool ctrl = (WinApi.GetAsyncKeyState(VK_CONTROL) & 0x8000) != 0;

        // 格式化为描述字符串
        string desc = ctrl ? $"Ctrl+{GetKeyName(vkCode)}" : GetKeyName(vkCode);

        // 调度到 UI 线程更新界面
        Application.Current?.Dispatcher.Invoke(() =>
        {
            RecentInputs.Insert(0, new InputEvent { Description = desc, ... });
        });
    }
    return WinApi.CallNextHookEx(hookId, nCode, wParam, lParam);
}
''')

add_heading(doc, '8.4 XAML — 声明式 UI（对比 tkinter.ttk）', 2, 14)
add_normal_para(doc, 'tkinter.ttk 的声明式写法：')
add_code_block(doc, '''\
# Python tkinter
root = Tk()
frame = Frame(root, bg="#F5F8FF", padx=10, pady=10)
frame.pack(fill=BOTH, expand=True)

label = Label(frame, text="输入监控", font=("微软雅黑", 13, "bold"))
label.pack(side=LEFT)

btn = Button(frame, text="📊", command=toggle_stats)
btn.pack(side=RIGHT)

listbox = Listbox(frame, bg="white", bd=0)
listbox.pack(fill=BOTH, expand=True)
''')

add_normal_para(doc, 'XAML 的声明式写法（更声明式，UI 和逻辑分离）：')
add_code_block(doc, '''\
<!-- MainWindow.xaml -->
<Window ... Title="输入显示器" Width="340" Height="540"
         WindowStyle="None" AllowsTransparency="True" Topmost="True">

    <!-- 资源字典（类似 CSS） -->
    <Window.Resources>
        <Style x:Key="HeaderButton" TargetType="Button">
            <Setter Property="Width" Value="28"/>
            <Setter Property="Background" Value="Transparent"/>
            <!-- ... -->
        </Style>
    </Window.Resources>

    <!-- 主面板（类似 tkinter 的 Frame） -->
    <Border Background="#F5F8FF" CornerRadius="12">

        <!-- 标题栏 -->
        <Border CornerRadius="12,12,0,0" Background="#E8EEFF">
            <Grid>
                <TextBlock Text="输入监控" FontSize="13" FontWeight="Bold"/>
                <Button Content="📊" Click="StatsToggle_Click"
                        HorizontalAlignment="Right"/>
            </Grid>
        </Border>

        <!-- 数据绑定：类似 tkinter 的 StringVar/IntVar 自动更新 -->
        <ListBox ItemsSource="{Binding RecentInputs}"
                 ItemTemplate="{StaticResource InputItemTemplate}"/>

    </Border>
</Window>
''')
add_normal_para(doc, 'XAML 中的 `{Binding ...}` 类似于 tkinter 的 `StringVar` 和 `IntVar`，但强大得多——支持路径、转换器、模式。')

# ════════════════════════════════════════════════════════════
# 第九章
# ════════════════════════════════════════════════════════════
add_heading(doc, '九、WPF 数据绑定：MVVM 思想', 1)

add_heading(doc, '9.1 什么是 MVVM', 2, 14)
add_normal_para(doc, 'MVVM（Model-View-ViewModel）是 WPF 推崇的架构模式，核心思想是：')
add_bullet(doc, '**Model（模型）**：数据本身（InputEvent、SessionStatistics）')
add_bullet(doc, '**View（视图）**：XAML 界面，只负责展示，不写业务逻辑')
add_bullet(doc, '**ViewModel（视图模型）**：连接 Model 和 View，处理业务逻辑，通过 INPC 通知 View 更新')
add_normal_para(doc, '对比 Python 的 MVC/MVP 模式：tkinter 中的大部分逻辑直接写在回调里，而 WPF 强制你把数据、界面、业务逻辑分开。')

add_heading(doc, '9.2 数据绑定的威力（对比 tkinter）', 2, 14)
add_normal_para(doc, 'tkinter 中更新 UI 需要手动调用（被动式）：')
add_code_block(doc, '''\
# Python tkinter - 手动更新（容易遗漏）
def update_label(count):
    label.config(text=f"计数: {count}")  # 必须手动更新

def on_click():
    global counter
    counter += 1
    update_label(counter)  # 容易忘记调用
''')

add_normal_para(doc, 'WPF 中数据绑定是自动的（主动式）：')
add_code_block(doc, '''\
<!-- XAML - 声明绑定，无需手动调用 -->
<TextBlock>
    <Run Text="键盘按键: "/>
    <Run Text="{Binding Statistics.TotalKeyPresses}"/>  <!-- 自动刷新 -->
</TextBlock>

<!-- C# - 只需修改属性值，UI 自动更新 -->
Statistics.TotalKeyPresses++;  // 触发 INPC，UI 自动刷新
// 不需要调用任何 update 方法
''')

add_heading(doc, '9.3 ObservableCollection（对比 list + 手动刷新）', 2, 14)
add_code_block(doc, '''\
# Python - 列表变了需要手动 notify
class MyModel:
    def __init__(self):
        self.items = []
        self.callbacks = []

    def add_item(self, item):
        self.items.insert(0, item)
        for cb in self.callbacks:  # 手动通知
            cb(self.items)

# tkinter 中需要 trace_write 或手动调用 listbox.insert()''')

add_code_block(doc, '''\
// C# - ObservableCollection 自动通知
public ObservableCollection<InputEvent> RecentInputs { get; } = new();

// 任何地方调用 Insert，ListBox 自动刷新（不需要 notify）
RecentInputs.Insert(0, inputEvent);
// 内部会触发 INotifyCollectionChanged
// WPF 的 ListBox 自动重新渲染
''')

# ════════════════════════════════════════════════════════════
# 第十章
# ════════════════════════════════════════════════════════════
add_heading(doc, '十、线程模型：协程 vs Dispatcher', 1)

add_heading(doc, '10.1 Python 异步模型', 2, 14)
add_code_block(doc, '''\
# Python 协程模型（单线程协作式多任务）
import asyncio

async def keyboard_reader():
    while True:
        event = await read_keyboard()    # 挂起，等待事件
        await process_event(event)      # 事件到达，处理
        await asyncio.sleep(0)          # 让出控制权

async def ui_updater():
    while True:
        await update_ui()               # 定期刷新界面
        await asyncio.sleep(0.016)      # ~60fps

asyncio.run(asyncio.gather(keyboard_reader(), ui_updater()))
# 全部在同一线程中运行，无需锁''')

add_heading(doc, '10.2 WPF / Win32 的线程模型', 2, 14)
add_code_block(doc, '''\
┌────────────────────────────────────────────────────────────┐
│                      UI 线程（主线程）                       │
│  • 运行 WPF Dispatcher 消息循环                            │
│  • 处理所有 UI 更新                                         │
│  • 所有 XAML 绑定、INPC 通知在此线程执行                    │
└─────────────────────────┬──────────────────────────────────┘
                          │ Dispatcher.Invoke() / BeginInvoke()
                          │ （Marshal 代码到 UI 线程）
┌─────────────────────────▼──────────────────────────────────┐
│                   钩子回调线程                              │
│  • CRT 的 Low-Level 钩子线程（系统线程）                   │
│  • 键盘/鼠标事件在此线程触发回调                          │
│  • 不能直接操作 UI（会崩溃或死锁）                        │
└───────────────────────────────────────────────────────────┘''')

add_normal_para(doc, '类比理解：WPF 的 UI 线程类似 Python asyncio 的主线程，'
    '但它不是用 await 切换协程，而是用 Dispatcher 将代码「投递」到 UI 线程中执行。')

add_heading(doc, '10.3 为什么需要 Dispatcher', 2, 14)
add_code_block(doc, '''\
# ❌ 错误：在钩子线程中直接操作 UI（类似在非主线程操作 tkinter）
def keyboard_callback(key):
    listbox.insert(0, key)  # tkinter: 不是线程安全的
    label.config(text=key) # WPF: 可能死锁、崩溃

# ✅ 正确：调度到 UI 线程
Application.Current?.Dispatcher.Invoke(() =>
{
    // 这里面的代码会在 UI 线程中执行
    RecentInputs.Insert(0, inputEvent);  // 安全
    label.Text = desc;                     // 安全
});

// 类似 Python 的 queue 或 asyncio 的 loop.call_soon_threadsafe
''')

add_heading(doc, '10.4 Invoke vs BeginInvoke', 2, 14)
add_code_block(doc, '''\
// Invoke：同步等待，等待 UI 线程执行完才返回
// （类似 asyncio 的 await）
Dispatcher.Invoke(() => { /* UI 更新 */ });

// BeginInvoke：异步投递，不等待，立即返回
// （类似 asyncio 的 create_task）
Dispatcher.BeginInvoke(() => { /* UI 更新 */ });

// 项目中为什么用 Invoke 而不是 BeginInvoke？
// 因为钩子回调是高频的（每次按键），BeginInvoke 可能导致消息堆积。
// Invoke 虽然阻塞当前线程（但那是钩子线程，不影响 UI），确保顺序性。
''')

# ════════════════════════════════════════════════════════════
# 第十一章
# ════════════════════════════════════════════════════════════
add_heading(doc, '十一、编译、打包与发布', 1)

add_heading(doc, '11.1 Debug vs Release 构建', 2, 14)
add_code_block(doc, '''\
# Debug 构建（开发用，包含调试符号）
dotnet build
# 输出：bin/Debug/net8.0-windows/InputDisplay.dll + InputDisplay.exe

# Release 构建（发布用，代码优化）
dotnet build -c Release
# 输出：bin/Release/net8.0-windows/InputDisplay.dll + InputDisplay.exe
''')

add_heading(doc, '11.2 发布（Publish）', 2, 14)
add_code_block(doc, '''\
# 框架依赖版（需要目标机器安装 .NET Runtime）
dotnet publish -c Release -r win-x64 --self-contained false -o ./publish

# 自包含版（不需要安装 .NET，文件约 60-80MB）
dotnet publish -c Release -r win-x64 --self-contained true -o ./publish-standalone

# 发布产物：
# InputDisplay.exe     ← 双击运行
# InputDisplay.dll     ← 托管程序集
# InputDisplay.deps.json
# InputDisplay.runtimeconfig.json
''')

add_heading(doc, '11.3 NuGet 包管理（对比 pip）', 2, 14)
add_code_block(doc, '''\
# Python: pip install xxx
pip install python-docx

# C# / NuGet: 在 .csproj 中声明，自动恢复
<ItemGroup>
    <PackageReference Include="CommunityToolkit.Mvvm" Version="8.2.2" />
</ItemGroup>

# 或者用命令行
dotnet add package CommunityToolkit.Mvvm

# NuGet 会下载包到 ~/.nuget/packages/，无需提交到仓库
# 团队成员 clone 后运行 dotnet restore 即可
''')

# ════════════════════════════════════════════════════════════
# 第十二章
# ════════════════════════════════════════════════════════════
add_heading(doc, '十二、学习路径与资源推荐', 1)

add_heading(doc, '12.1 如果你想深入学习 WPF', 2, 14)
add_bullet(doc, '官方文档：https://learn.microsoft.com/zh-cn/dotnet/desktop/wpf/')
add_bullet(doc, '书籍：《WPF 编程宝典》（Pro WPF in C#）')
add_bullet(doc, '视频：Bilibili 搜索"WPF 入门教程"')

add_heading(doc, '12.2 如果你想用 Python 实现类似功能', 2, 14)
add_bullet(doc, 'pynput：https://pypi.org/project/pynput/ — 键盘鼠标监听')
add_bullet(doc, 'PyQt / PySide：跨平台 GUI，支持透明窗口、数据绑定')
add_bullet(doc, 'ctypes：调用 Win32 API 的标准库方案')
add_bullet(doc, 'pywin32：更高级的 Windows API 封装')

add_heading(doc, '12.3 如果你想用 Python 实现 WPF 类似的声明式 GUI', 2, 14)
add_bullet(doc, 'NiceGUI：https://nicegui.io — Python 优先的现代 Web 风格 GUI')
add_bullet(doc, 'CustomTkinter：美化版的 tkinter')
add_bullet(doc, 'Flet：用 Flutter 构建 Python GUI，支持数据绑定')

add_heading(doc, '12.4 项目代码阅读顺序建议', 2, 14)
add_bullet(doc, '第一步：读 WinApi.cs — 理解底层 API 封装')
add_bullet(doc, '第二步：读 Models.cs — 理解数据模型和 INPC')
add_bullet(doc, '第三步：读 InputHookService.cs — 理解钩子回调和数据处理')
add_bullet(doc, '第四步：读 MainWindow.xaml — 理解 UI 布局和数据绑定')
add_bullet(doc, '第五步：读 MainWindow.xaml.cs — 理解交互逻辑和线程协调')

add_divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— 文档结束 —')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── 保存 ────────────────────────────────────────────────────
output_path = 'F:/aicodingtools/claudecode_project/project4/github-upload/docs/InputDisplay_Project_Guide_for_Python_Developers.docx'
doc.save(output_path)
print(f'Done! Word document saved to: {output_path}')
